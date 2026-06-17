import os
import pdfplumber
from docx import Document

def extract_text(file_path: str) -> str:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext == ".docx":
        return _extract_from_docx(file_path)
    else:
        raise ValueError(
            "Unsupported file format. Only PDF and DOCX are supported."
        )

def _extract_from_pdf(file_path: str) -> str:
    pages = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text.strip())
        full_text = "\n".join(pages)
        if not full_text.strip():
            raise RuntimeError("No text found in PDF.")
        return full_text
    except Exception as e:
        raise RuntimeError(
            f"Failed to extract PDF text: {e}"
        )

def _extract_from_docx(file_path: str) -> str:
    lines = []
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        lines.append(text)
        full_text = "\n".join(lines)
        if not full_text.strip():
            raise RuntimeError("No text found in DOCX.")
        return full_text
    except Exception as e:
        raise RuntimeError(
            f"Failed to extract DOCX text: {e}"
        )

def get_file_info(file_path: str) -> dict:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    return {
        "file_name": os.path.basename(file_path),
        "file_type": os.path.splitext(file_path)[1].lower(),
        "file_size_kb": round(os.path.getsize(file_path) / 1024, 2)
    }

